import datetime
import os
import re
import subprocess
import urllib.parse
import uuid
from io import BytesIO
from typing import Optional
import redis
import boto3
import bson
from PyPDF2 import PdfReader
from bson import ObjectId
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx2pdf import convert
from dotenv import dotenv_values
from fastapi import APIRouter, status, Form, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pymongo import MongoClient
from models import SearchlightAPIModel

redis_client = redis.Redis(host='localhost', port=6000, db=0)
router = APIRouter()
pdf_text_list = []
config = dotenv_values(".env")
mongodb_client = MongoClient(config["ATLAS_URI"])
mongodb_database = mongodb_client[config["DB_NAME"]]
pipeline_request = mongodb_database[config["COLLECTION_NAME"]]
raw_request = mongodb_database[config["RAW_COLLECTION_NAME"]]

AWS_ACCESS_KEY = config["AWS_ACCESS_KEY"]
AWS_SECRET_KEY = config["AWS_SECRET_KEY"]
BUCKET_NAME = config["BUCKET_NAME"]
FOLDER_NAME = 'uploads'

def cache_key(keyword: str, pdf_filename: str) -> str:
    return f"{keyword}:{pdf_filename}"


def upload_file_to_s3(file, bucket_name, folder_name, object_key):
    s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

    try:
        s3_client.put_object(
            Bucket=bucket_name,
            Key=f'{folder_name}/{object_key}',
            Body=file,
            ACL='public-read',
        )
    except Exception as e:
        (
            print("Error:", e))
    object_url = f"https://{bucket_name}.s3.ap-south-1.amazonaws.com/{folder_name}/{object_key}"
    print(f"File '{object_key}' in folder '{folder_name}' is now publicly accessible at: {object_url}")
    return object_url


def pdf_to_text(pdf_path):
    read_pdf = PdfReader(pdf_path)
    pdf_pages = len(read_pdf.pages)
    all_text = []
    for page_number in range(pdf_pages):
        pdf_page = read_pdf.pages[page_number]
        text_extracted_from_pdf = pdf_page.extract_text()
        all_text.append(text_extracted_from_pdf)
    return all_text


def word_count(text, search_word):
    pattern = re.compile(r'\b' + re.escape(search_word) + r'\b', re.IGNORECASE)
    matches = pattern.findall(text)
    count = len(matches)
    return count


def txt_to_doc(input_text, output_file, search_word):
    doc = Document()
    file_content = re.sub(r"[^\x00-\x7F]+|\x0c", " ", input_text)
    paragraphs = file_content.split('\n')
    pattern = re.compile(r'\b' + re.escape(search_word) + r'\b', re.IGNORECASE)
    for paragraph in paragraphs:
        p = doc.add_paragraph()
        matches = list(pattern.finditer(paragraph))
        start = 0
        for match in matches:
            p.add_run(paragraph[start:match.start()])
            run = p.add_run(match.group())
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            start = match.end()
        p.add_run(paragraph[start:])
    doc.save(output_file)


def doc_to_pdf(doc_file, pdf_output):
    # convert(doc_file, pdf_output)
    pdf_file_path = pdf_output + '/' + doc_file.replace('.docx', '.pdf')
    try:
        subprocess.call([
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', pdf_output,
            doc_file
        ])
        return pdf_file_path
    except Exception as e:
        return str(e)


@router.post("/upload", response_description="SearchLight API",
             status_code=status.HTTP_201_CREATED,
             response_model=JudgmentAPIModel
             )
async def upload_pdf(
        keyword: str = Form(...),
        pdf: UploadFile = File(...),
):
    key = cache_key(keyword, pdf.filename)
    cached_response = redis_client.get(key)
    if cached_response:
        return JSONResponse(
            content=cached_response,
            headers={"Cache-Control": "private, max-age=3600"}
        )
    uploaded_pdf = pdf.file.read()
    text_in_pdf = pdf_to_text(BytesIO(uploaded_pdf))
    pdf_text_list.extend(text_in_pdf)
    text_extracted = '\n'.join(pdf_text_list)
    count = word_count(text_extracted, keyword)
    txt_to_doc(text_extracted, 'txt-to-doc.docx', keyword)
    guid = uuid.uuid4().hex
    OBJECT_KEY = guid + '.pdf'
    doc_to_pdf('txt-to-doc.docx', OBJECT_KEY)
    with open(guid + '.pdf', 'rb') as generated_pdf_file:
        generated_pdf_content = generated_pdf_file.read()
    object_url = upload_file_to_s3(generated_pdf_content, BUCKET_NAME, FOLDER_NAME, OBJECT_KEY)

    api_data = {
        "keyword": keyword,
        "docName": pdf.filename,
        "docType": pdf.content_type,
        "keywordCount": count,
        "docUrl": object_url
    }

    redis_client.set(key, api_data, ex=3600)
    return JSONResponse(content=api_data,headers={"Cache-Control": "private, max-age=3600"})